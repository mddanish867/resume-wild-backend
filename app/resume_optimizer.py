import os
import re
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import pipeline, DistilBertTokenizer, DistilBertForMaskedLM
from sklearn.feature_extraction.text import CountVectorizer, ENGLISH_STOP_WORDS
from nltk.tokenize import sent_tokenize
from collections import Counter
import nltk

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load models with error handling
try:
    MODEL_NAME = "distilbert-base-uncased"
    tokenizer = DistilBertTokenizer.from_pretrained(MODEL_NAME)
    model = DistilBertForMaskedLM.from_pretrained(MODEL_NAME)
    fill_mask = pipeline(task="fill-mask", model=model, tokenizer=tokenizer)
    logger.info("BERT model loaded successfully")
except Exception as e:
    logger.warning(f"Could not load BERT model: {e}")
    fill_mask = None

class ResumeOptimizer:
    def __init__(self):
        self.processed_keywords = set()
        self.keyword_density_limit = 0.03  # 3% max density per section
        self.keywords_added_count = 0
        
    def extract_keywords(self, text, top_k=30):
        """Extract meaningful keywords with better filtering"""
        if not text or len(text.strip()) < 10:
            return []
            
        # Clean text first
        text = self.clean_text(text)
        
        # Extended stop words for resume context
        custom_stop_words = list(ENGLISH_STOP_WORDS) + [
            'experience', 'work', 'skills', 'projects', 'job', 'role', 'position',
            'company', 'team', 'using', 'worked', 'including', 'various', 'multiple',
            'different', 'new', 'good', 'strong', 'excellent', 'responsible', 'duties',
            'tasks', 'requirements', 'qualifications', 'candidate', 'ideal', 'looking',
            'seeking', 'years', 'year', 'month', 'months', 'time', 'day', 'days'
        ]
        
        vectorizer = CountVectorizer(
            ngram_range=(1, 3),  # Include 3-grams for better context
            stop_words=custom_stop_words,
            max_features=top_k * 3,
            min_df=1,
            token_pattern=r'\b[a-zA-Z][a-zA-Z0-9+#.\-]{1,}\b'  # Better token pattern
        )
        
        try:
            matrix = vectorizer.fit_transform([text])
            keywords = list(zip(vectorizer.get_feature_names_out(), matrix.sum(axis=0).tolist()[0]))
            
            # Filter and sort keywords
            filtered_keywords = []
            for word, score in sorted(keywords, key=lambda x: x[1], reverse=True):
                if self.is_valid_keyword(word):
                    filtered_keywords.append(word)
                    
            return filtered_keywords[:top_k]
        except Exception as e:
            logger.warning(f"Keyword extraction failed: {e}")
            return []

    def is_valid_keyword(self, word):
        """Check if a keyword is valid for resume optimization"""
        if len(word) < 2:
            return False
        if word.isdigit():
            return False
        if word.lower() in ['http', 'https', 'www', 'com', 'org', 'net']:
            return False
        # Allow technical terms with numbers and special characters
        if re.match(r'^[a-zA-Z][a-zA-Z0-9+#.\-]*[a-zA-Z0-9+#]$', word):
            return True
        if re.match(r'^[a-zA-Z]{2,}$', word):
            return True
        return False

    def clean_text(self, text):
        """Enhanced text cleaning"""
        if not text:
            return ""
        
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text)
        # Keep alphanumeric, spaces, and common technical characters
        text = re.sub(r'[^\w\s+#.\-]', ' ', text)
        return text.strip()

    def get_missing_keywords(self, resume_text, jd_text, max_keywords=20):
        """Get missing keywords with better filtering"""
        resume_keywords = set([kw.lower() for kw in self.extract_keywords(resume_text)])
        jd_keywords = self.extract_keywords(jd_text, top_k=50)
        
        missing = []
        for kw in jd_keywords:
            if (kw.lower() not in resume_keywords and 
                len(kw) > 2 and 
                kw.lower() not in self.processed_keywords and
                self.is_technical_or_skill_keyword(kw)):
                missing.append(kw)
                
        return missing[:max_keywords]

    def is_technical_or_skill_keyword(self, keyword):
        """Check if keyword is likely a technical skill or relevant term"""
        technical_indicators = [
            # Programming languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 'go',
            # Technologies
            'react', 'angular', 'vue', 'node', 'express', 'django', 'flask', 'spring',
            # Databases
            'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle',
            # Cloud
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins',
            # Methodologies
            'agile', 'scrum', 'devops', 'ci/cd', 'tdd', 'bdd',
            # General skills
            'api', 'rest', 'graphql', 'microservices', 'testing', 'debugging'
        ]
        
        keyword_lower = keyword.lower()
        
        # Check if it's a known technical term
        if any(tech in keyword_lower for tech in technical_indicators):
            return True
            
        # Check if it looks like a technical term (has version numbers, etc.)
        if re.search(r'[0-9]+\.[0-9]+|\d+', keyword):
            return True
            
        # Check if it's an acronym (2-5 uppercase letters)
        if re.match(r'^[A-Z]{2,5}$', keyword):
            return True
            
        # Check if it's a compound technical term
        if len(keyword.split()) > 1 and any(word.lower() in technical_indicators for word in keyword.split()):
            return True
            
        return len(keyword) > 3  # Default: accept longer terms

    def check_keyword_density(self, text, keyword):
        """Check if adding keyword would exceed density limit"""
        words = text.split()
        if len(words) < 10:
            return True
            
        keyword_count = text.lower().count(keyword.lower())
        density = keyword_count / len(words)
        return density < self.keyword_density_limit

    def enhance_sentence_contextually(self, sentence, keyword, section_type):
        """Add keywords contextually based on section type"""
        if not self.check_keyword_density(sentence, keyword):
            return sentence, False
            
        sentence = sentence.strip()
        if not sentence:
            return sentence, False
            
        # Avoid duplicates
        if keyword.lower() in sentence.lower():
            return sentence, False
            
        # Section-specific enhancements
        templates = {
            'skills': [
                f"Proficient in {keyword}",
                f"Experienced with {keyword}",
                f"Skilled in {keyword}",
                f"Expert in {keyword}"
            ],
            'experience': [
                f"Utilized {keyword} for development",
                f"Implemented solutions using {keyword}",
                f"Worked extensively with {keyword}",
                f"Applied {keyword} in professional projects",
                f"Leveraged {keyword} technology"
            ],
            'projects': [
                f"Built using {keyword}",
                f"Implemented with {keyword}",
                f"Developed using {keyword}",
                f"Created solutions with {keyword}"
            ],
            'summary': [
                f"Experienced in {keyword}",
                f"Proficient in {keyword}",
                f"Skilled professional with {keyword} expertise",
                f"Strong background in {keyword}"
            ]
        }
        
        section_templates = templates.get(section_type, templates['experience'])
        
        # Choose template based on context
        template_index = min(len(section_templates) - 1, len(sentence.split()) // 20)
        enhancement = f" {section_templates[template_index]}."
            
        # Mark keyword as processed
        self.processed_keywords.add(keyword.lower())
        self.keywords_added_count += 1
        
        return sentence + enhancement, True

    def preserve_document_formatting(self, source_doc, target_doc):
        """Preserve original document formatting"""
        try:
            # Copy document properties
            if hasattr(source_doc, 'core_properties') and hasattr(target_doc, 'core_properties'):
                target_doc.core_properties.author = source_doc.core_properties.author
                target_doc.core_properties.title = source_doc.core_properties.title
            
            # Copy sections if they exist
            if hasattr(source_doc, 'sections') and hasattr(target_doc, 'sections'):
                for i, section in enumerate(source_doc.sections):
                    if i < len(target_doc.sections):
                        target_section = target_doc.sections[i]
                        try:
                            target_section.page_height = section.page_height
                            target_section.page_width = section.page_width
                            target_section.left_margin = section.left_margin
                            target_section.right_margin = section.right_margin
                            target_section.top_margin = section.top_margin
                            target_section.bottom_margin = section.bottom_margin
                        except Exception as e:
                            logger.warning(f"Could not copy section formatting: {e}")
                        
        except Exception as e:
            logger.warning(f"Could not preserve document formatting: {e}")

    def copy_paragraph_formatting(self, source_para, target_para):
        """Copy paragraph formatting from source to target"""
        try:
            # Copy paragraph alignment
            target_para.alignment = source_para.alignment
            
            # Copy paragraph formatting
            if hasattr(source_para, 'paragraph_format') and hasattr(target_para, 'paragraph_format'):
                pf_source = source_para.paragraph_format
                pf_target = target_para.paragraph_format
                
                if pf_source.space_before:
                    pf_target.space_before = pf_source.space_before
                if pf_source.space_after:
                    pf_target.space_after = pf_source.space_after
                if pf_source.line_spacing:
                    pf_target.line_spacing = pf_source.line_spacing
                if pf_source.first_line_indent:
                    pf_target.first_line_indent = pf_source.first_line_indent
                if pf_source.left_indent:
                    pf_target.left_indent = pf_source.left_indent
                if pf_source.right_indent:
                    pf_target.right_indent = pf_source.right_indent
            
            # Copy run formatting for each run in the paragraph
            for i, run in enumerate(source_para.runs):
                if i < len(target_para.runs):
                    target_run = target_para.runs[i]
                    try:
                        target_run.bold = run.bold
                        target_run.italic = run.italic
                        target_run.underline = run.underline
                        if run.font.name:
                            target_run.font.name = run.font.name
                        if run.font.size:
                            target_run.font.size = run.font.size
                    except Exception as re:
                        logger.warning(f"Could not copy run formatting: {re}")
                        
        except Exception as e:
            logger.warning(f"Could not copy paragraph formatting: {e}")

    def detect_section_type(self, text):
        """Detect section type based on content"""
        if not text:
            return 'other'
            
        text_lower = text.lower().strip()
        
        # Check if it's a section header (short text with section keywords)
        section_headers = {
            'summary': ['summary', 'profile', 'objective', 'about me', 'professional summary'],
            'skills': ['skills', 'technical skills', 'technologies', 'competencies', 'expertise'],
            'experience': ['experience', 'work experience', 'employment', 'work history', 'professional experience'],
            'projects': ['projects', 'portfolio', 'work samples', 'key projects'],
            'education': ['education', 'academic background', 'degree', 'university', 'college'],
            'awards': ['awards', 'achievements', 'honors', 'recognition', 'certifications']
        }
        
        # First check for exact section header matches
        for section_type, keywords in section_headers.items():
            if any(text_lower == keyword or text_lower.startswith(keyword) for keyword in keywords):
                return section_type
        
        # Then check for section keywords in content
        for section_type, keywords in section_headers.items():
            if any(keyword in text_lower for keyword in keywords):
                return section_type
                
        # Content-based classification
        if '|' in text and any(tech in text_lower for tech in ['python', 'java', 'javascript', 'sql', 'react']):
            return 'skills'
        elif any(word in text_lower for word in ['developed', 'implemented', 'created', 'built', 'designed']):
            return 'projects'
        elif any(word in text_lower for word in ['worked', 'managed', 'led', 'responsible', 'collaborated']):
            return 'experience'
        elif any(word in text_lower for word in ['degree', 'university', 'college', 'graduated']):
            return 'education'
            
        return 'other'

    def is_section_header(self, text):
        """Check if text is likely a section header"""
        if not text:
            return False
            
        text_clean = text.strip()
        words = text_clean.split()
        
        # Headers are usually short (1-4 words)
        if len(words) > 4:
            return False
            
        # Check for common header patterns
        header_keywords = [
            'summary', 'profile', 'objective', 'skills', 'experience', 
            'projects', 'education', 'awards', 'certifications', 'achievements'
        ]
        
        return any(keyword in text.lower() for keyword in header_keywords)

    def remove_duplicates(self, text):
        """Remove duplicate sentences and phrases"""
        if not text:
            return text
            
        try:
            sentences = sent_tokenize(text)
        except:
            # Fallback if NLTK fails
            sentences = [s.strip() for s in text.split('.') if s.strip()]
            
        unique_sentences = []
        seen = set()
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Normalize sentence for comparison
            normalized = re.sub(r'\s+', ' ', sentence.lower())
            normalized = re.sub(r'[^\w\s]', '', normalized)
            
            if normalized not in seen and len(sentence) > 5:
                unique_sentences.append(sentence)
                seen.add(normalized)
                
        return ' '.join(unique_sentences)

    def optimize_skills_section(self, text, missing_keywords, max_additions=8):
        """Optimize skills section specifically"""
        if '|' in text:
            # Pipe-separated skills
            current_skills = [skill.strip() for skill in text.split('|')]
            relevant_keywords = [kw for kw in missing_keywords if kw not in text.lower()][:max_additions]
            
            if relevant_keywords:
                updated_skills = current_skills + relevant_keywords
                return ' | '.join(updated_skills), len(relevant_keywords)
        else:
            # Comma-separated or other format
            relevant_keywords = [kw for kw in missing_keywords if kw not in text.lower()][:max_additions]
            if relevant_keywords:
                addition = ' | ' + ' | '.join(relevant_keywords)
                return text + addition, len(relevant_keywords)
                
        return text, 0

    def optimize_resume_docx(self, input_path, jd_text, output_path):
        """Main optimization function with improved logic"""
        try:
            if not os.path.exists(input_path):
                raise FileNotFoundError(f"Input file not found: {input_path}")
                
            # Clean and prepare job description
            jd_text_clean = self.clean_text(jd_text)
            if len(jd_text_clean) < 20:
                raise ValueError("Job description too short or invalid")
            
            # Read original resume content
            source_doc = Document(input_path)
            resume_content = ' '.join([p.text for p in source_doc.paragraphs if p.text.strip()])
            
            if len(resume_content) < 50:
                raise ValueError("Resume content too short or invalid")
            
            # Get missing keywords
            missing_keywords = self.get_missing_keywords(resume_content, jd_text_clean)
            logger.info(f"Found {len(missing_keywords)} missing keywords: {missing_keywords[:10]}")
            
            if not missing_keywords:
                logger.info("No missing keywords found, copying original document")
                # Just copy the original if no optimization needed
                source_doc.save(output_path)
                return 0
            
            # Create new document
            optimized_doc = Document()
            self.preserve_document_formatting(source_doc, optimized_doc)
            
            # Reset counters for this document
            self.processed_keywords.clear()
            self.keywords_added_count = 0
            
            # Section-specific keyword limits
            max_keywords_per_section = {
                'skills': 8, 
                'experience': 5, 
                'projects': 4, 
                'summary': 3, 
                'other': 2
            }
            
            current_section = 'other'
            section_keywords_used = 0
            
            # Process each paragraph
            for para_idx, para in enumerate(source_doc.paragraphs):
                text = para.text.strip()
                
                if not text:
                    # Add empty paragraph to maintain structure
                    optimized_doc.add_paragraph("")
                    continue
                
                # Detect section type
                detected_section = self.detect_section_type(text)
                
                # Check if this is a new section header
                if self.is_section_header(text):
                    current_section = detected_section
                    section_keywords_used = 0
                    # Keep headers as-is
                    new_para = optimized_doc.add_paragraph(text)
                    self.copy_paragraph_formatting(para, new_para)
                    logger.info(f"Section header detected: {current_section}")
                    continue
                
                # Process content paragraphs
                enhanced_text = text
                keywords_added_to_para = 0
                max_for_section = max_keywords_per_section.get(current_section, 2)
                
                # Special handling for skills section
                if current_section == 'skills' and section_keywords_used < max_for_section:
                    remaining_budget = max_for_section - section_keywords_used
                    enhanced_text, added_count = self.optimize_skills_section(
                        text, missing_keywords, remaining_budget
                    )
                    keywords_added_to_para = added_count
                    section_keywords_used += added_count
                
                # Handle other sections
                elif current_section in ['experience', 'projects', 'summary'] and section_keywords_used < max_for_section:
                    available_keywords = [kw for kw in missing_keywords if kw.lower() not in enhanced_text.lower()]
                    
                    for keyword in available_keywords:
                        if (section_keywords_used < max_for_section and 
                            self.keywords_added_count < 15):  # Global limit
                            
                            result, added = self.enhance_sentence_contextually(
                                enhanced_text, keyword, current_section
                            )
                            
                            if added:
                                enhanced_text = result
                                keywords_added_to_para += 1
                                section_keywords_used += 1
                
                # Clean up the enhanced text
                enhanced_text = self.remove_duplicates(enhanced_text)
                enhanced_text = self.clean_text(enhanced_text)
                
                # Add optimized paragraph
                new_para = optimized_doc.add_paragraph(enhanced_text)
                self.copy_paragraph_formatting(para, new_para)
                
                if keywords_added_to_para > 0:
                    logger.info(f"Added {keywords_added_to_para} keywords to {current_section} section")
            
            # Save optimized document
            optimized_doc.save(output_path)
            logger.info(f"Optimized resume saved to {output_path} with {self.keywords_added_count} keywords added")
            
            return self.keywords_added_count
            
        except Exception as e:
            logger.error(f"Resume optimization failed: {e}")
            raise

# Create global instance for backward compatibility
resume_optimizer = ResumeOptimizer()

def optimize_resume_docx(input_path, jd_text, output_path):
    """Wrapper function for backward compatibility"""
    return resume_optimizer.optimize_resume_docx(input_path, jd_text, output_path)

def read_paragraphs(path):
    """Helper function to read paragraphs from docx"""
    try:
        doc = Document(path)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        logger.error(f"Failed to read paragraphs from {path}: {e}")
        return []

# Example usage
if __name__ == "__main__":
    optimizer = ResumeOptimizer()
    resume_path = "uploads/sample_resume.docx"
    jd_text = """
    We are looking for a .NET full stack developer skilled in Azure Functions, DevOps, CI/CD,
    React, SQL, cloud deployments, and writing scalable APIs. Experience with microservices,
    Docker, Kubernetes, and Agile methodologies is preferred.
    """
    output_path = "optimized/optimized_resume.docx"
    
    try:
        keywords_added = optimizer.optimize_resume_docx(resume_path, jd_text, output_path)
        print(f"Resume optimized successfully with {keywords_added} keywords added")
    except Exception as e:
        print(f"Optimization failed: {e}")