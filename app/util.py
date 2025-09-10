import os
import docx
import logging
import re
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

def extract_sections_from_docx(path: str) -> Dict[str, List[str]]:
    """
    Extract content from .docx and return section-wise content with improved detection.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")

    try:
        doc = docx.Document(path)
        sections = {
            "summary": [],
            "skills": [],
            "experience": [],
            "projects": [],
            "education": [],
            "awards": [],
            "certifications": [],
            "other": []
        }

        current_section = "other"
        section_keywords = {
            "summary": ["summary", "profile", "objective", "about", "professional summary"],
            "skills": ["skills", "technical", "technologies", "competencies", "expertise"],
            "experience": ["experience", "employment", "work history", "professional experience", "career"],
            "projects": ["projects", "portfolio", "work samples", "key projects"],
            "education": ["education", "academic", "degree", "university", "college"],
            "awards": ["awards", "achievements", "honors", "recognition"],
            "certifications": ["certifications", "certificates", "licensed"]
        }

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            lower_text = text.lower()
            
            # Check if this paragraph is a section header
            is_section_header = False
            for section_name, keywords in section_keywords.items():
                if any(keyword in lower_text and len(text.split()) <= 5 for keyword in keywords):
                    current_section = section_name
                    is_section_header = True
                    break
            
            # Add content to appropriate section
            if not is_section_header:
                sections[current_section].append(text)
            else:
                # Also add the header itself
                sections[current_section].append(text)

        # Clean empty sections
        cleaned_sections = {k: v for k, v in sections.items() if v}
        
        logger.info(f"Extracted sections: {list(cleaned_sections.keys())}")
        return cleaned_sections

    except Exception as e:
        logger.error(f"Failed to extract from DOCX: {e}")
        raise

def extract_text_from_docx(path: str) -> str:
    """
    Extract all text content from a DOCX file.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    
    try:
        doc = docx.Document(path)
        full_text = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        
        return '\n'.join(full_text)
    
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise

def validate_docx_file(path: str) -> tuple[bool, str]:
    """
    Validate if a DOCX file is readable and contains content.
    Returns (is_valid, message)
    """
    try:
        if not os.path.exists(path):
            return False, f"File not found: {path}"
        
        if not path.lower().endswith('.docx'):
            return False, "File is not a DOCX file"
        
        # Check file size
        file_size = os.path.getsize(path)
        if file_size == 0:
            return False, "File is empty"
        
        if file_size > 10 * 1024 * 1024:  # 10MB limit
            return False, "File too large (max 10MB)"
        
        # Try to read the document
        doc = docx.Document(path)
        
        # Check if document has content
        has_content = False
        for para in doc.paragraphs:
            if para.text.strip():
                has_content = True
                break
        
        if not has_content:
            return False, "Document appears to be empty"
        
        return True, "Valid DOCX file"
    
    except Exception as e:
        return False, f"Error reading DOCX file: {str(e)}"

def clean_filename(filename: str) -> str:
    """
    Clean filename to remove invalid characters.
    """
    # Remove or replace invalid characters
    cleaned = re.sub(r'[<>:"/\\|?*]', '_', filename)
    
    # Remove multiple consecutive underscores
    cleaned = re.sub(r'_+', '_', cleaned)
    
    # Remove leading/trailing underscores and dots
    cleaned = cleaned.strip('_.')
    
    # Ensure filename is not too long
    name, ext = os.path.splitext(cleaned)
    if len(name) > 100:
        name = name[:100]
    
    return name + ext

def get_file_info(path: str) -> Dict:
    """
    Get comprehensive file information.
    """
    if not os.path.exists(path):
        return {"exists": False}
    
    stat = os.stat(path)
    
    return {
        "exists": True,
        "size": stat.st_size,
        "size_mb": round(stat.st_size / (1024 * 1024), 2),
        "modified": stat.st_mtime,
        "is_readable": os.access(path, os.R_OK),
        "is_writable": os.access(path, os.W_OK),
        "extension": os.path.splitext(path)[1].lower()
    }

def extract_keywords_from_text(text: str, min_length: int = 2, max_keywords: int = 50) -> List[str]:
    """
    Extract potential keywords from text using simple heuristics.
    """
    if not text:
        return []
    
    # Split text into words
    words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#.\-]*\b', text.lower())
    
    # Common stop words to exclude
    stop_words = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
        'by', 'from', 'up', 'about', 'into', 'through', 'during', 'before', 'after',
        'above', 'below', 'between', 'among', 'is', 'are', 'was', 'were', 'be', 'been',
        'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
        'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those', 'i', 'you',
        'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them', 'my', 'your',
        'his', 'her', 'its', 'our', 'their'
    }
    
    # Filter words
    keywords = []
    for word in words:
        if (len(word) >= min_length and 
            word not in stop_words and
            not word.isdigit()):
            keywords.append(word)
    
    # Count frequency and return most common
    from collections import Counter
    word_counts = Counter(keywords)
    
    return [word for word, count in word_counts.most_common(max_keywords)]

def format_file_size(size_bytes: int) -> str:
    """
    Format file size in human readable format.
    """
    if size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB"]
    i = 0
    size = float(size_bytes)
    
    while size >= 1024.0 and i < len(size_names) - 1:
        size /= 1024.0
        i += 1
    
    return f"{size:.1f} {size_names[i]}"

def create_backup_filename(original_path: str) -> str:
    """
    Create a backup filename by adding timestamp.
    """
    from datetime import datetime
    
    name, ext = os.path.splitext(original_path)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    return f"{name}_backup_{timestamp}{ext}"

def ensure_directory_structure():
    """
    Ensure all required directories exist for the application.
    """
    required_dirs = [
        "uploads",
        "optimized", 
        "backups",
        "logs",
        "instance"
    ]
    
    created_dirs = []
    
    for directory in required_dirs:
        if not os.path.exists(directory):
            try:
                os.makedirs(directory, exist_ok=True)
                created_dirs.append(directory)
                logger.info(f"Created directory: {directory}")
            except Exception as e:
                logger.error(f"Failed to create directory {directory}: {e}")
                
    return created_dirs

def log_file_operation(operation: str, file_path: str, success: bool, details: str = ""):
    """
    Log file operations for debugging and monitoring.
    """
    status = "SUCCESS" if success else "FAILED"
    message = f"FILE_OP: {operation} | {file_path} | {status}"
    
    if details:
        message += f" | {details}"
    
    if success:
        logger.info(message)
    else:
        logger.error(message)

def safe_file_delete(file_path: str) -> bool:
    """
    Safely delete a file with error handling.
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            log_file_operation("DELETE", file_path, True)
            return True
        else:
            logger.warning(f"File not found for deletion: {file_path}")
            return False
    except Exception as e:
        log_file_operation("DELETE", file_path, False, str(e))
        return False

def get_document_stats(path: str) -> Dict:
    """
    Get statistics about a DOCX document.
    """
    if not os.path.exists(path):
        return {"error": "File not found"}
    
    try:
        doc = docx.Document(path)
        
        paragraphs = doc.paragraphs
        total_paragraphs = len(paragraphs)
        non_empty_paragraphs = len([p for p in paragraphs if p.text.strip()])
        
        # Count words
        total_words = 0
        total_chars = 0
        
        for para in paragraphs:
            text = para.text.strip()
            if text:
                words = len(text.split())
                total_words += words
                total_chars += len(text)
        
        # Detect sections
        sections = extract_sections_from_docx(path)
        section_count = len([s for s in sections.values() if s])
        
        return {
            "total_paragraphs": total_paragraphs,
            "non_empty_paragraphs": non_empty_paragraphs,
            "total_words": total_words,
            "total_characters": total_chars,
            "sections_detected": section_count,
            "sections": list(sections.keys()),
            "avg_words_per_paragraph": round(total_words / non_empty_paragraphs, 2) if non_empty_paragraphs > 0 else 0
        }
        
    except Exception as e:
        logger.error(f"Failed to get document stats: {e}")
        return {"error": str(e)}

def compare_documents(original_path: str, optimized_path: str) -> Dict:
    """
    Compare original and optimized documents to show changes.
    """
    try:
        original_stats = get_document_stats(original_path)
        optimized_stats = get_document_stats(optimized_path)
        
        if "error" in original_stats or "error" in optimized_stats:
            return {"error": "Could not analyze one or both documents"}
        
        word_diff = optimized_stats["total_words"] - original_stats["total_words"]
        char_diff = optimized_stats["total_characters"] - original_stats["total_characters"]
        para_diff = optimized_stats["non_empty_paragraphs"] - original_stats["non_empty_paragraphs"]
        
        return {
            "original": original_stats,
            "optimized": optimized_stats,
            "differences": {
                "words_added": word_diff,
                "characters_added": char_diff,
                "paragraphs_added": para_diff,
                "percentage_increase": round((word_diff / original_stats["total_words"]) * 100, 2) if original_stats["total_words"] > 0 else 0
            }
        }
        
    except Exception as e:
        logger.error(f"Document comparison failed: {e}")
        return {"error": str(e)}

def cleanup_old_files(directory: str, max_age_days: int = 7) -> List[str]:
    """
    Clean up files older than specified days in a directory.
    """
    import time
    
    if not os.path.exists(directory):
        return []
    
    current_time = time.time()
    max_age_seconds = max_age_days * 24 * 60 * 60
    deleted_files = []
    
    try:
        for filename in os.listdir(directory):
            file_path = os.path.join(directory, filename)
            
            if os.path.isfile(file_path):
                file_age = current_time - os.path.getmtime(file_path)
                
                if file_age > max_age_seconds:
                    if safe_file_delete(file_path):
                        deleted_files.append(filename)
    
    except Exception as e:
        logger.error(f"Cleanup failed for directory {directory}: {e}")
    
    if deleted_files:
        logger.info(f"Cleaned up {len(deleted_files)} old files from {directory}")
    
    return deleted_files

def validate_job_description(job_desc: str) -> tuple[bool, str]:
    """
    Validate job description content for optimization.
    """
    if not job_desc or not isinstance(job_desc, str):
        return False, "Job description is required"
    
    job_desc = job_desc.strip()
    
    if len(job_desc) < 50:
        return False, "Job description too short (minimum 50 characters)"
    
    if len(job_desc) > 5000:
        return False, "Job description too long (maximum 5000 characters)"
    
    # Check for meaningful content
    words = job_desc.split()
    if len(words) < 10:
        return False, "Job description should contain at least 10 words"
    
    # Check for common job-related keywords
    job_keywords = [
        'experience', 'skills', 'requirements', 'responsibilities', 'qualifications',
        'developer', 'engineer', 'manager', 'analyst', 'specialist', 'coordinator',
        'required', 'preferred', 'must', 'should', 'knowledge', 'proficient'
    ]
    
    has_job_keywords = any(keyword in job_desc.lower() for keyword in job_keywords)
    if not has_job_keywords:
        return False, "Job description should contain job-related keywords"
    
    return True, "Valid job description"

# Utility function for testing
def create_sample_resume(output_path: str) -> bool:
    """
    Create a sample resume for testing purposes.
    """
    try:
        doc = docx.Document()
        
        # Add title
        title = doc.add_paragraph()
        title_run = title.runs[0] if title.runs else title.add_run()
        title_run.text = "John Doe"
        title_run.bold = True
        title_run.font.size = docx.shared.Pt(16)
        title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact info
        contact = doc.add_paragraph("john.doe@email.com | (555) 123-4567 | LinkedIn: linkedin.com/in/johndoe")
        contact.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Add summary
        doc.add_paragraph("\nPROFESSIONAL SUMMARY")
        doc.add_paragraph("Experienced software developer with 5+ years in full-stack development. Proven track record in building scalable web applications and leading development teams.")
        
        # Add skills
        doc.add_paragraph("\nTECHNICAL SKILLS")
        doc.add_paragraph("Programming Languages: Python | Java | JavaScript")
        doc.add_paragraph("Frameworks: Django | React | Spring Boot")
        doc.add_paragraph("Databases: MySQL | PostgreSQL | MongoDB")
        
        # Add experience
        doc.add_paragraph("\nWORK EXPERIENCE")
        doc.add_paragraph("Senior Software Developer | TechCorp Inc. | 2020-Present")
        doc.add_paragraph("• Developed and maintained web applications using Python and React")
        doc.add_paragraph("• Led a team of 4 developers in agile development processes")
        doc.add_paragraph("• Implemented CI/CD pipelines reducing deployment time by 40%")
        
        doc.add_paragraph("\nSoftware Developer | StartupXYZ | 2018-2020")
        doc.add_paragraph("• Built RESTful APIs using Django and PostgreSQL")
        doc.add_paragraph("• Collaborated with cross-functional teams to deliver features")
        
        # Add education
        doc.add_paragraph("\nEDUCATION")
        doc.add_paragraph("Bachelor of Science in Computer Science | University of Technology | 2018")
        
        # Save document
        doc.save(output_path)
        logger.info(f"Sample resume created: {output_path}")
        return True
        
    except Exception as e:
        logger.error(f"Failed to create sample resume: {e}")
        return False